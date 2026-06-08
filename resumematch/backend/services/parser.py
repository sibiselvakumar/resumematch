import io
import pdfplumber
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using pdfplumber."""
    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {e}")
    return "\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {e}")


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract text from a file based on its extension."""
    filename_lower = filename.lower()
    if filename_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif filename_lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif filename_lower.endswith(".doc"):
        raise ValueError("Legacy .doc format is not supported. Please convert to .docx or .pdf.")
    else:
        # Plain text fallback
        return file_bytes.decode("utf-8", errors="ignore")


def get_candidate_name(filename: str) -> str:
    """Derive a candidate display name from the filename."""
    name = filename.rsplit(".", 1)[0]  # Remove extension
    name = name.replace("_", " ").replace("-", " ")
    # Remove common resume prefixes like "CV", "Resume"
    for prefix in ["cv ", "resume ", "resume_", "cv_"]:
        if name.lower().startswith(prefix):
            name = name[len(prefix):]
    return name.strip().title() or filename
